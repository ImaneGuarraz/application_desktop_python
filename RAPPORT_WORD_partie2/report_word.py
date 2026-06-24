# report_word.py

"""
Génération du document Word pour le rapport Frankenstein
"""

from docx import Document   # crée le document Word
from docx.shared import Inches   # gère les tailles d'images


def create_report_docx(metadata: dict, stats: dict, chart_path: str, image_path: str, output_docx: str):
    doc = Document()   # crée un nouveau document

    title = metadata["title"]   # récupère le titre
    author = metadata["author"]   # récupère l'auteur

    doc.add_heading(title, level=0)   # titre principal
    doc.add_paragraph(f"Auteur du livre : {author}")   # auteur du livre
    doc.add_paragraph("Auteur du rapport : Imane")   # ton nom

    doc.add_picture(image_path, width=Inches(5.5))   # ajoute l'image finale
    doc.add_page_break()   # nouvelle page

    doc.add_heading("Distribution des longueurs des paragraphes", level=1)   # titre page 2
    doc.add_picture(chart_path, width=Inches(5.5))   # ajoute le graphique

    doc.add_heading("Description", level=2)   # sous-titre

    # Résumé d’intrigue (obligatoire dans l’énoncé)
    doc.add_paragraph(
        "Résumé de l'intrigue : Frankenstein raconte l’histoire de Victor Frankenstein, "
        "un jeune scientifique qui crée une créature vivante à partir de morceaux de corps humains. "
        "Effrayé par sa propre création, il l’abandonne, ce qui déclenche une série de tragédies "
        "où la créature cherche désespérément compréhension et vengeance."   # résumé original
    )

    # Statistiques demandées
    doc.add_paragraph(f"Nombre total de paragraphes : {stats['nb_paragraphs']}")   # stats
    doc.add_paragraph(f"Nombre total de mots : {stats['total_words']}")   # stats
    doc.add_paragraph(f"Paragraphe le plus court : {stats['min_words']} mots")   # stats
    doc.add_paragraph(f"Paragraphe le plus long : {stats['max_words']} mots")   # stats
    doc.add_paragraph(f"Moyenne de mots par paragraphe : {stats['avg_words']:.2f}")   # stats
    doc.add_paragraph("Source : Project Gutenberg")   # source

    doc.save(output_docx)   # enregistre le document
    return output_docx   # retourne le chemin
